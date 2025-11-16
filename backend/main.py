from fastapi import FastAPI, Form, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from openai import OpenAI
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import io
import json
from datetime import datetime
import time
from PyPDF2 import PdfReader

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

# === JSEARCH API ===
JSEARCH_URL = "https://jsearch.p.rapidapi.com/search"
JSEARCH_HEADERS = {
    "X-RapidAPI-Key": "18f96b9949msh2eb8bd4f154081ap10b6b9jsnac95d87233c3",
    "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
}

client = OpenAI(api_key="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")

preferred_jobs = []

# ===================== CV GENERATOR - USER PROVIDES FIELD (NO HARDCODE) =====================
@app.post("/generate-cv")
async def generate_cv(
    full_name: str = Form(""), 
    email: str = Form(""), 
    contact: str = Form(""), 
    address: str = Form(""),
    field: str = Form(""),  # NEW: USER INPUT FOR FIELD
    summary: str = Form(""),
    exp1_role: str = Form(""), exp1_company: str = Form(""), exp1_duration: str = Form(""), exp1_description: str = Form(""),
    exp2_role: str = Form(""), exp2_company: str = Form(""), exp2_duration: str = Form(""), exp2_description: str = Form(""),
    exp3_role: str = Form(""), exp3_company: str = Form(""), exp3_duration: str = Form(""), exp3_description: str = Form(""),
    exp4_role: str = Form(""), exp4_company: str = Form(""), exp4_duration: str = Form(""), exp4_description: str = Form(""),
    exp5_role: str = Form(""), exp5_company: str = Form(""), exp5_duration: str = Form(""), exp5_description: str = Form(""),
    projects: str = Form(""), 
    skills: str = Form(""),
    edu1_degree: str = Form(""), edu1_university: str = Form(""), edu1_duration: str = Form(""),
    edu2_degree: str = Form(""), edu2_university: str = Form(""), edu2_duration: str = Form(""),
    edu3_degree: str = Form(""), edu3_university: str = Form(""), edu3_duration: str = Form(""),
    edu4_degree: str = Form(""), edu4_university: str = Form(""), edu4_duration: str = Form(""),
    edu5_degree: str = Form(""), edu5_university: str = Form(""), edu5_duration: str = Form("")
):
    doc = docx.Document()

    # NAME
    p = doc.add_paragraph()
    run = p.add_run(full_name.upper())
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # FIELD (USER INPUT — NO HARDCODE)
    p = doc.add_paragraph()
    run = p.add_run(field.upper())  # ← USER KA INPUT YAHAN
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = "Calibri"

    # CONTACT
    p = doc.add_paragraph(f"{address} | {email} | {contact}")
    p.paragraph_format.space_after = Pt(12)

    # LINE
    p = doc.add_paragraph()
    p.add_run("─" * 60)

    # SUMMARY
    doc.add_heading("SUMMARY", level=2)
    doc.add_paragraph(summary)

    # EXPERIENCE
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
    for i in range(1, 6):
        role = locals().get(f"exp{i}_role", "")
        company = locals().get(f"exp{i}_company", "")
        duration = locals().get(f"exp{i}_duration", "")
        desc = locals().get(f"exp{i}_description", "")
        if role or company:
            p = doc.add_paragraph()
            p.add_run(role).bold = True
            p.add_run(" " * 70 + duration)
            if desc:
                for line in [l.strip() for l in desc.split('\n') if l.strip()]:
                    doc.add_paragraph(f"• {line}", style="List Bullet")

    # PROJECTS
    doc.add_heading("PROJECTS", level=2)
    for line in projects.split("\n"):
        if line.strip():
            doc.add_paragraph(f"• {line.strip()}", style="List Bullet")

    # SKILLS
    doc.add_heading("SKILLS", level=2)
    skills_list = [s.strip() for s in skills.split(",") if s.strip()]
    rows = (len(skills_list) + 2) // 3
    table = doc.add_table(rows=rows, cols=3)
    table.style = "Table Grid"
    table.allow_autofit = False
    for col in table.columns:
        col.width = Inches(2.3)
    for i, skill in enumerate(skills_list):
        cell = table.cell(i // 3, i % 3)
        cell.text = skill
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # EDUCATION
    doc.add_heading("EDUCATION", level=2)
    for i in range(1, 6):
        degree = locals().get(f"edu{i}_degree", "")
        university = locals().get(f"edu{i}_university", "")
        duration = locals().get(f"edu{i}_duration", "")
        if degree or university:
            p = doc.add_paragraph()
            p.add_run(degree).bold = True
            p.add_run(" " * 70 + duration)
            if university:
                doc.add_paragraph(university)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={full_name.replace(' ', '_')}_CV.docx"}
    )

# ===================== AI-ONLY SCORING + IMPROVEMENTS =====================
@app.post("/match-jobs")
async def match_jobs(cv_file: UploadFile = File(...), keyword: str = Form(""), posted_time: str = Form("all")):
    content = await cv_file.read()
    if cv_file.filename.lower().endswith('.pdf'):
        cv_text = " ".join([p.extract_text() or "" for p in PdfReader(io.BytesIO(content)).pages])
    else:
        cv_text = " ".join([p.text for p in docx.Document(io.BytesIO(content)).paragraphs])
    cv_text = cv_text[:4000]

    time.sleep(2)

    query_params = {
        "query": f"{keyword} remote" if "remote" not in keyword.lower() else keyword,
        "page": "1",
        "num_pages": "1"
    }

    try:
        response = requests.get(JSEARCH_URL, headers=JSEARCH_HEADERS, params=query_params, timeout=15)
        data = response.json()
        jobs = data.get("data", [])[:10]
    except:
        return {"matches": [], "message": "JSearch API error"}

    now = datetime.utcnow()
    matches = []

    for i, job in enumerate(jobs):
        job_title = job.get("job_title", "Unknown")
        company = job.get("employer_name", "Unknown")
        location = job.get("job_city", "Remote")
        salary = job.get("job_salary", "Not disclosed")
        apply_link = job.get("job_apply_link", "#")
        posted_date = job.get("job_posted_at_datetime_utc", "")

        try:
            job_dt = datetime.fromisoformat(posted_date.replace("Z", "+00:00"))
            if posted_time == "24h" and (now - job_dt).days > 1: continue
            if posted_time == "7d" and (now - job_dt).days > 7: continue
        except: pass

        job_desc = job.get("job_description", "")
        req_skills = ", ".join(job.get("job_required_skills", []))

        prompt = f"""
You are a senior hiring manager. Analyze the CV and job.

CV:
{cv_text}

Job Title: {job_title}
Company: {company}
Description: {job_desc}
Required Skills: {req_skills}

Give:
1. A score from 0 to 10 (how well the candidate matches the job).
2. 3 short, specific improvements to reach 9–10/10.

Return ONLY JSON:
{{
  "score": 8,
  "improvements": ["Add React experience", "Show leadership", "Tailor summary"]
}}
"""

        try:
            res = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                max_tokens=200
            )
            result = json.loads(res.choices[0].message.content)
            score = int(result.get("score", 0))
            improvements_raw = result.get("improvements", [])
            improvements = "\n".join([f"• {item}".strip() for item in improvements_raw if item])
        except:
            score = 0
            improvements = "• AI error\n• Try again"

        matches.append({
            "id": str(i),
            "job": {
                "title": job_title,
                "company": company,
                "logo": job.get("employer_logo") or "https://via.placeholder.com/80",
                "location": location,
                "salary": salary,
                "apply_link": apply_link
            },
            "score": score,
            "improvements": improvements
        })

    return {"matches": sorted(matches, key=lambda x: x["score"], reverse=True)}

# ===================== PREFER ENDPOINTS =====================
@app.post("/prefer-job")
async def prefer_job(job_data: dict):
    preferred_jobs.append(job_data)
    return {"message": "Job preferred!"}

@app.get("/preferred-jobs")
async def get_preferred():
    return {"preferred": preferred_jobs}